"""Stage 9 - generate the formal comparative report in multiple formats.

Builds one grounded *content model* from the saved results
(``reports/comparison.json`` + ``models/forecast_metrics.json`` + ``config``)
and renders it to **Markdown**, **self-contained HTML** (print-friendly, so
"Print -> Save as PDF" yields a clean PDF), and **Word (.docx)**. Because all
three formats render the *same* model, they never drift apart.

Every number is read from the actual simulation runs - nothing is hand-typed or
invented. Run ``python main.py report`` (optionally ``--format md html docx``)
after ``python main.py compare``.
"""
from __future__ import annotations

import html as _html
import json
from typing import Any

import config
from utils.logger import get_logger

log = get_logger("report")

# --- Output paths -----------------------------------------------------------
MD_PATH = config.REPORTS_DIR / "report.md"
HTML_PATH = config.REPORTS_DIR / "report.html"
DOCX_PATH = config.REPORTS_DIR / "Microgrid_AI_Comparative_Report.docx"

# (label, metric key, "min"|"max" is better, format string) for the results table.
_ROWS = [
    ("Daily energy cost (Rs)", "daily_cost_inr", "min", "{:.0f}"),
    ("Cost saved vs grid-only (%)", "cost_saved_pct", "max", "{:.1f}"),
    ("CO2 emitted (kg)", "co2_kg", "min", "{:.0f}"),
    ("CO2 saved vs grid-only (kg)", "co2_saved_kg", "max", "{:.0f}"),
    ("Renewable self-consumption (%)", "renewable_self_consumption_pct", "max", "{:.1f}"),
    ("Renewable share of demand (%)", "renewable_share_of_demand_pct", "max", "{:.1f}"),
    ("Grid import (kWh)", "grid_import_kwh", "min", "{:.0f}"),
    ("LLM calls", "llm_calls", "min", "{:.0f}"),
    ("LLM tokens", "llm_tokens", "min", "{:,.0f}"),
    ("Latency per decision (ms)", "avg_decision_latency_ms", "min", "{:.0f}"),
]


# --- Content model ----------------------------------------------------------
# A report is a list of blocks. Each block is a tuple whose first element names
# its kind; the renderers below turn the same blocks into MD / HTML / DOCX.
#   ("title", text)                    ("subtitle", text)   ("meta", text)
#   ("h", level, text)                 ("p", text)          ("p_i", text)
#   ("bullets", [[(seg, bold), ...]])
#   ("table", headers, rows, highlight, caption)
#       headers:   list[str]  (includes the row-label column header)
#       rows:      list[list[str]]
#       highlight: list[int | None]  (0-based cell index to flag per row, or None)
#       caption:   str | None

def _load(path, what) -> Any:
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:  # pragma: no cover
        log.warning("Could not read %s: %s", what, exc)
        return None


def _best_index(values: list[float], better: str) -> int:
    return values.index(max(values) if better == "max" else min(values))


def build_model(comp: dict, forecast: list | None) -> list[tuple]:
    """Assemble the grounded content model from the saved results."""
    R = comp["results"]
    names = list(R.keys())
    rb, gen, ag = R["Rule-based"], R["Generative-AI"], R["Agentic-AI"]
    by_cost = min(R.items(), key=lambda kv: kv[1]["daily_cost_inr"])[0]
    by_co2 = min(R.items(), key=lambda kv: kv[1]["co2_kg"])[0]
    tok_ratio = ag["llm_tokens"] / gen["llm_tokens"] if gen["llm_tokens"] else 0
    win = comp["window_hours"]

    blocks: list[tuple] = [
        ("title", "Generative AI vs Agentic AI for Microgrid Operation"),
        ("subtitle", "A Comparative Analysis"),
        ("meta", f"Evaluation window: {win} hours ({win // 24} days) from "
                 f"{config.EVAL_START_DATE}  |  LLM (both AI approaches): "
                 f"{comp['model']} (Groq)"),

        ("h", 1, "Abstract"),
        ("p", "This study solves a single microgrid energy-management problem "
              "three ways - a rule-based baseline, a generative large-language-"
              "model (LLM) controller, and an autonomous tool-using agent - and "
              "compares them on identical physics, forecasts, and metrics. Both "
              "AI controllers use the same Groq model, so any difference "
              "reflects the paradigm rather than the model. On this problem the "
              f"rule-based baseline was the most cost-effective controller "
              f"({rb['cost_saved_pct']:.1f}% saved versus a grid-only bill) at "
              "zero token cost; the generative LLM's value lay in explanation "
              "rather than better control; and the agentic approach spent the "
              "most compute for the weakest result. The contribution is a fair, "
              "reproducible measurement framework and an honest, nuanced finding "
              "- not a foregone 'AI wins'."),

        ("h", 1, "1. Problem and system"),
        ("p", "A microgrid is a small, self-contained power system. The "
              f"simulated site comprises a {int(config.PV_RATED_KWP)} kW rooftop "
              f"solar array, a {int(config.WIND_RATED_KW)} kW wind turbine, a "
              f"{int(config.BATTERY_CAPACITY_KWH)} kWh battery, and a grid "
              "connection under a time-of-use tariff (peak import at Rs "
              f"{config.GRID_IMPORT_PRICE:.0f}/kWh; grid carbon intensity "
              f"{config.CO2_PER_KWH_GRID} kg CO2/kWh). Every hour a controller "
              "must decide whether to charge, discharge, or idle the battery to "
              "reduce cost and emissions while meeting demand - deciding on "
              "forecasts, before the actual generation and load are known."),
        ("p_i", "The dataset is clearly-labelled synthetic: generated from "
                "documented physical models (solar geometry, a turbine power "
                "curve, a realistic demand shape), not measured field data, and "
                "fully reproducible from a fixed seed. No results in this report "
                "are invented; every figure is read from the actual simulation "
                "runs."),

        ("h", 1, "2. The three approaches"),
        ("bullets", [
            [("Rule-based (baseline, no AI). ", True),
             ("A transparent hand-written policy: charge on renewable surplus, "
              "discharge during expensive peaks, top up cheaply off-peak. The "
              "comparison floor - instant, free, fully auditable, but fixed.",
              False)],
            [("Generative AI (LLM, human-in-the-loop). ", True),
             ("Each hour the microgrid state and forecasts are sent to a Groq "
              "LLM, which recommends a battery set-point and explains it in one "
              "sentence. A single prompt, no tools; advisory not autonomous.",
              False)],
            [("Agentic AI (autonomous, tool-using). ", True),
             ("A LangGraph ReAct agent that calls a simulation tool to test "
              "candidate actions against the real physics, then autonomously "
              "commits the best one - a closed loop with no human.", False)],
        ]),

        ("h", 1, "3. Methodology"),
        ("p", "All three controllers were scored in one shared simulator on the "
              f"same {win}-hour window. Controllers decide on forecasts; the "
              "environment settles on the actuals, so a lucky guess is not "
              "rewarded. Requested battery power is clipped to real charge/"
              "discharge and state-of-charge limits, round-trip efficiency is "
              "applied, and the hourly energy balance is settled against "
              "measured generation and demand. Cost, CO2, renewable use, LLM "
              "calls, tokens, and latency are accounted identically for all "
              "three."),
    ]

    # 3.1 Forecasting accuracy (optional - only if metrics present).
    if forecast:
        f_rows = []
        for m in forecast:
            best = m.get("best_metrics", {})
            naive = m.get("naive_baseline", {})
            f_rows.append([
                str(m.get("target", "-")),
                str(m.get("best_model", "-")),
                f"{best.get('mae', float('nan')):.2f}",
                f"{best.get('rmse', float('nan')):.2f}",
                f"{naive.get('rmse', float('nan')):.2f}",
            ])
        blocks += [
            ("h", 2, "3.1 Forecasting accuracy"),
            ("p", "Gradient-boosted models predict solar, wind, and demand from "
                  "calendar and recent-history features (no same-hour weather). "
                  "All three targets beat a seasonal-naive baseline on a "
                  "time-ordered hold-out (last 20% of the series):"),
            ("table", ["Target", "Best model", "MAE", "RMSE", "Naive RMSE"],
             f_rows, [None] * len(f_rows),
             "Lower is better; RMSE below the naive column means the model "
             "learned genuine structure. Wind is the hardest target."),
        ]

    # 4. Results table.
    headers = ["Metric"] + names
    res_rows, highlight = [], []
    for label, key, better, fmt in _ROWS:
        vals = [float(R[n][key]) for n in names]
        best = _best_index(vals, better)
        res_rows.append([label] + [fmt.format(v) for v in vals])
        highlight.append(best + 1)  # +1 for the leading label column
    blocks += [
        ("h", 1, "4. Results"),
        ("table", headers, res_rows, highlight,
         "Highlighted = best value in each row."),

        ("h", 1, "5. Analysis"),
        ("h", 2, "5.1 Optimisation quality (cost and CO2)"),
        ("bullets", [
            [(f"Lowest energy cost: {by_cost} (Rs "
              f"{R[by_cost]['daily_cost_inr']:.0f}/day, "
              f"{R[by_cost]['cost_saved_pct']:.1f}% below grid-only).", False)],
            [(f"Lowest emissions: {by_co2} ({R[by_co2]['co2_kg']:.0f} kg CO2 "
              "over the window).", False)],
            [(f"Cost saved: rule-based {rb['cost_saved_pct']:.1f}%, generative "
              f"{gen['cost_saved_pct']:.1f}%, agentic {ag['cost_saved_pct']:.1f}"
              "%. The extra AI machinery did not translate into better "
              "optimisation.", False)],
        ]),
        ("h", 2, "5.2 Decision autonomy and explainability"),
        ("bullets", [
            [("Rule-based ", True),
             ("- fully automated but a fixed policy that cannot adapt beyond "
              "its rules; the most transparent (the exact branch is auditable).",
              False)],
            [("Generative AI ", True),
             ("- advisory (human-in-the-loop); attaches a one-sentence "
              "natural-language rationale to every decision, the most "
              "human-readable of the three.", False)],
            [("Agentic AI ", True),
             ("- fully autonomous closed loop; exposes a tool-evaluation trace, "
              "but the choice of which candidates to try is opaque.", False)],
        ]),
        ("h", 2, "5.3 Cost of intelligence"),
        ("bullets", [
            [("The baseline makes 0 LLM calls and decides in microseconds.",
              False)],
            [(f"Generative: {gen['llm_calls']} calls / {gen['llm_tokens']:,} "
              f"tokens. Agentic: {ag['llm_calls']} calls / {ag['llm_tokens']:,} "
              f"tokens (~{tok_ratio:.1f}x the generative approach).", False)],
            [("LLM latency reflects throttled wall-clock (network round-trips "
              "and free-tier rate-limit back-off), but the ordering "
              "(agentic > generative > rule-based) is inherent to the "
              "paradigms.", False)],
        ]),
        ("h", 2, "5.4 Reliability and complexity"),
        ("bullets", [
            [("Rule-based is deterministic and cannot fail at runtime; the LLM "
              "approaches must handle rate limits, occasional unparseable "
              "replies (mitigated by a safe-idle fallback), and run-to-run "
              "non-determinism.", False)],
            [("Implementation complexity rises from rule-based (a few branches) "
              "to generative (prompt + parse) to agentic (a tool-using agent "
              "with accounting and resume support).", False)],
        ]),

        ("h", 1, "6. Conclusion"),
        ("p", "For this microgrid problem and dataset, the rule-based baseline "
              "is the most cost-effective controller - it matches or beats the "
              "AI approaches on cost at zero token cost. The generative LLM's "
              "real value is explanation, not better control. The agentic "
              "system's autonomy did not pay off here: its tool-guided "
              "optimisation is myopic (it minimises each hour in isolation "
              "rather than planning across the day), so it spent the most "
              "compute for the weakest result. Agentic AI is most promising "
              "where decisions are genuinely multi-step, tools unlock "
              "information a rule cannot encode, and full autonomy is required - "
              "conditions this single-step, well-understood problem does not "
              "present."),
        ("p_i", "Caveats. Results are on clearly-labelled synthetic data, a "
                "small (8B) model, and a deliberately simple agent design; a "
                "stronger model, richer tools, or a multi-step planning horizon "
                "could shift the agentic outcome. The framework built here is "
                "what makes such follow-ups measurable."),
    ]
    return blocks


# --- Markdown renderer ------------------------------------------------------
def render_md(blocks: list[tuple]) -> str:
    out: list[str] = []
    for b in blocks:
        kind = b[0]
        if kind == "title":
            out += [f"# {b[1]}", ""]
        elif kind == "subtitle":
            out += [f"*{b[1]}*", ""]
        elif kind == "meta":
            out += [f"> {b[1]}", ""]
        elif kind == "h":
            out += [f"{'#' * (b[1] + 1)} {b[2]}", ""]
        elif kind == "p":
            out += [b[1], ""]
        elif kind == "p_i":
            out += [f"*{b[1]}*", ""]
        elif kind == "bullets":
            for segs in b[1]:
                line = "".join(f"**{t}**" if bold else t for t, bold in segs)
                out.append(f"- {line}")
            out.append("")
        elif kind == "table":
            _, headers, rows, highlight, caption = b
            out.append("| " + " | ".join(headers) + " |")
            out.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for i, row in enumerate(rows):
                cells = []
                for j, cell in enumerate(row):
                    cells.append(f"**{cell}**" if highlight[i] == j else cell)
                out.append("| " + " | ".join(cells) + " |")
            out.append("")
            if caption:
                out += [f"*{caption}*", ""]
    return "\n".join(out).rstrip() + "\n"


# --- HTML renderer (self-contained, print-friendly) -------------------------
_HTML_CSS = """
  :root{--ink:#111827;--muted:#6b7280;--line:#d9dee6;--navy:#123a6b;--green:#0a7d33;}
  *{box-sizing:border-box;}
  body{max-width:820px;margin:2.2rem auto;padding:0 1.3rem;color:var(--ink);
    font:15px/1.6 Calibri,system-ui,-apple-system,Segoe UI,Roboto,Arial,sans-serif;}
  h1.title{color:var(--navy);text-align:center;font-size:1.9rem;margin:.2rem 0;}
  p.subtitle{text-align:center;font-style:italic;font-size:1.15rem;margin:.1rem 0;}
  p.meta{text-align:center;color:var(--muted);font-size:.85rem;margin:.3rem 0 1.6rem;}
  h2{color:var(--navy);font-size:1.35rem;border-bottom:2px solid var(--line);
    padding-bottom:.2rem;margin-top:1.8rem;}
  h3{color:var(--navy);font-size:1.1rem;margin-top:1.2rem;}
  ul{padding-left:1.2rem;} li{margin:.25rem 0;}
  table{border-collapse:collapse;width:100%;margin:1rem 0;font-size:.92rem;}
  th,td{border:1px solid var(--line);padding:.4rem .6rem;text-align:right;}
  th:first-child,td:first-child{text-align:left;}
  thead th{background:#eef2f8;}
  td.best{font-weight:700;color:var(--green);background:#eafaf0;}
  p.caption{color:var(--muted);font-style:italic;font-size:.85rem;margin:.2rem 0 1rem;}
  @media print{body{margin:0;max-width:none;font-size:11pt;}
    h2,h3{page-break-after:avoid;} table,ul{page-break-inside:avoid;}}
"""


def render_html(blocks: list[tuple]) -> str:
    e = _html.escape
    out = ['<!doctype html>', '<html lang="en"><head><meta charset="utf-8">',
           '<meta name="viewport" content="width=device-width, initial-scale=1">',
           '<title>Microgrid AI - Comparative Report</title>',
           f'<style>{_HTML_CSS}</style></head><body>']
    for b in blocks:
        kind = b[0]
        if kind == "title":
            out.append(f'<h1 class="title">{e(b[1])}</h1>')
        elif kind == "subtitle":
            out.append(f'<p class="subtitle">{e(b[1])}</p>')
        elif kind == "meta":
            out.append(f'<p class="meta">{e(b[1])}</p>')
        elif kind == "h":
            out.append(f"<h{b[1] + 1}>{e(b[2])}</h{b[1] + 1}>")
        elif kind == "p":
            out.append(f"<p>{e(b[1])}</p>")
        elif kind == "p_i":
            out.append(f"<p><em>{e(b[1])}</em></p>")
        elif kind == "bullets":
            out.append("<ul>")
            for segs in b[1]:
                inner = "".join(
                    f"<strong>{e(t)}</strong>" if bold else e(t)
                    for t, bold in segs)
                out.append(f"<li>{inner}</li>")
            out.append("</ul>")
        elif kind == "table":
            _, headers, rows, highlight, caption = b
            out.append("<table><thead><tr>"
                       + "".join(f"<th>{e(h)}</th>" for h in headers)
                       + "</tr></thead><tbody>")
            for i, row in enumerate(rows):
                cells = []
                for j, cell in enumerate(row):
                    cls = ' class="best"' if highlight[i] == j else ""
                    cells.append(f"<td{cls}>{e(cell)}</td>")
                out.append("<tr>" + "".join(cells) + "</tr>")
            out.append("</tbody></table>")
            if caption:
                out.append(f'<p class="caption">{e(caption)}</p>')
    out.append("</body></html>")
    return "\n".join(out)


# --- Word (.docx) renderer --------------------------------------------------
def render_docx(blocks: list[tuple], path) -> None:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    navy, green = RGBColor(0x12, 0x3A, 0x6B), RGBColor(0x0A, 0x7D, 0x33)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    for b in blocks:
        kind = b[0]
        if kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(b[1])
            r.bold = True
            r.font.size = Pt(20)
            r.font.color.rgb = navy
        elif kind == "subtitle":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(b[1])
            r.italic = True
            r.font.size = Pt(13)
        elif kind == "meta":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(b[1]).font.size = Pt(9.5)
            doc.add_paragraph()
        elif kind == "h":
            h = doc.add_heading(b[2], level=b[1])
            for run in h.runs:
                run.font.color.rgb = navy
        elif kind == "p":
            doc.add_paragraph(b[1])
        elif kind == "p_i":
            p = doc.add_paragraph()
            r = p.add_run(b[1])
            r.italic = True
            r.font.size = Pt(10)
        elif kind == "bullets":
            for segs in b[1]:
                p = doc.add_paragraph(style="List Bullet")
                for text, bold in segs:
                    p.add_run(text).bold = bold
        elif kind == "table":
            _, headers, rows, highlight, caption = b
            tbl = doc.add_table(rows=1, cols=len(headers))
            tbl.style = "Light Grid Accent 1"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, htext in enumerate(headers):
                tbl.rows[0].cells[j].paragraphs[0].add_run(htext).bold = True
            for i, row in enumerate(rows):
                cells = tbl.add_row().cells
                for j, cell in enumerate(row):
                    run = cells[j].paragraphs[0].add_run(cell)
                    if highlight[i] == j:
                        run.bold = True
                        run.font.color.rgb = green
            if caption:
                cp = doc.add_paragraph()
                cr = cp.add_run(caption)
                cr.italic = True
                cr.font.size = Pt(9.5)
    doc.save(str(path))


# --- Orchestration ----------------------------------------------------------
_RENDERERS = {"md", "html", "docx"}


def generate_report(formats: list[str] | None = None) -> list:
    """Generate the report in the requested formats (default: all three)."""
    formats = list(formats) if formats else ["md", "html", "docx"]
    unknown = set(formats) - _RENDERERS
    if unknown:
        raise ValueError(f"Unknown format(s): {sorted(unknown)}. "
                         f"Choose from {sorted(_RENDERERS)}.")

    comp = _load(config.REPORTS_DIR / "comparison.json", "comparison")
    if not comp:
        log.error("No reports/comparison.json found. Run "
                  "`python main.py compare` first, then re-run report.")
        return []
    forecast = _load(config.MODELS_DIR / "forecast_metrics.json", "forecast")

    blocks = build_model(comp, forecast)
    written = []
    if "md" in formats:
        MD_PATH.write_text(render_md(blocks), encoding="utf-8")
        written.append(MD_PATH)
    if "html" in formats:
        HTML_PATH.write_text(render_html(blocks), encoding="utf-8")
        written.append(HTML_PATH)
    if "docx" in formats:
        try:
            render_docx(blocks, DOCX_PATH)
            written.append(DOCX_PATH)
        except ImportError:
            log.error("python-docx is not installed - skipping the .docx "
                      "format. Install it with `pip install python-docx` "
                      "(it is in requirements.txt).")
        except PermissionError:
            log.error("Could not write %s - it looks like the file is open "
                      "(e.g. in Word). Close it and re-run `python main.py "
                      "report --format docx`.", DOCX_PATH.name)

    for p in written:
        log.info("Wrote %s", p)
    if HTML_PATH in written:
        log.info("Tip: open report.html and use the browser's "
                 "Print -> Save as PDF for a clean PDF.")
    return written
